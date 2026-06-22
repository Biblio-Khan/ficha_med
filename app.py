import streamlit as st
import pandas as pd
import requests
import io
import hashlib
import base64
from datetime import datetime, timedelta, timezone
from docx import Document
from docx.shared import Pt, Inches
from deep_translator import GoogleTranslator
from docx.enum.text import WD_ALIGN_PARAGRAPH

# =====================================================================
# CONFIGURAÇÕES DA PÁGINA E SESSÃO
# =====================================================================
st.set_page_config(page_title="BiblioKhan Médicas", page_icon="bibliokhan.ico", layout="wide")

# Definição do Fuso Horário de Brasília (UTC-3)
FUSO_BR = timezone(timedelta(hours=-3))

st.markdown("""
    <style>
    .stButton>button {
        background-color: #9B5DE5 !important;
        color: white !important;
        border-radius: 8px !important;
        border: none !important;
    }
    .stButton>button:hover {
        background-color: #8446CE !important;
        color: white !important;
    }
    .stAlert {
        border-left-color: #9B5DE5 !important;
    }
    </style>
""", unsafe_allow_html=True)

if 'autenticado' not in st.session_state: st.session_state.autenticado = False
if 'user_nome' not in st.session_state: st.session_state.user_nome = ""
if 'user_email' not in st.session_state: st.session_state.user_email = ""
if 'user_perfil' not in st.session_state: st.session_state.user_perfil = ""
if 'user_creditos' not in st.session_state: st.session_state.user_creditos = 0

if 'lista_assuntos' not in st.session_state: st.session_state.lista_assuntos = []
if 'autores' not in st.session_state: st.session_state.autores = [""]
if 'colaboradores' not in st.session_state: st.session_state.colaboradores = []
if 'mesh_limite' not in st.session_state: st.session_state.mesh_limite = 5
if 'ultimo_termo' not in st.session_state: st.session_state.ultimo_termo = ""
if 'fichas_lote' not in st.session_state: st.session_state.fichas_lote = []

# =====================================================================
# CONFIGURAÇÃO DA API DO GOOGLE (VIA STREAMLIT SECRETS)
# =====================================================================
try:
    URL_API_GOOGLE = st.secrets["URL_API_GOOGLE"]
except KeyError:
    st.error("Erro: A variável 'URL_API_GOOGLE' não foi configurada nos Secrets do Streamlit.")
    st.stop()

# =====================================================================
# FUNÇÕES DO SISTEMA DE CRÉDITOS E AUTENTICAÇÃO
# =====================================================================
def hash_senha(senha):
    return hashlib.sha256(senha.encode()).hexdigest()

def validar_credenciais(email, senha):
    payload = {"action": "login", "email": email.strip().lower(), "senha": hash_senha(senha)}
    try:
        resp = requests.post(URL_API_GOOGLE, json=payload, timeout=15)
        res_json = resp.json()
        if res_json.get("success", False):
            return True, res_json.get("nome"), res_json.get("perfil"), res_json.get("creditos", 0)
        return False, None, None, 0
    except:
        return False, None, None, 0

def enviar_notificacao_telegram(nome, email, pacote):
    """ Envia um alerta ao administrador via Telegram quando uma compra é feita """
    token = st.secrets.get("TELEGRAM_BOT_TOKEN")
    chat_id = st.secrets.get("TELEGRAM_ADMIN_CHAT_ID")
    if token and chat_id:
        msg = f"🔔 *Nova Compra de Créditos!*\n\n👤 *Usuário:* {nome}\n📧 *E-mail:* {email}\n📦 *Pacote:* {pacote} Fichas\n\n📌 _Verifique o Painel Admin para realizar a homologação._"
        try:
            requests.post(f"https://api.telegram.org/bot{token}/sendMessage", json={"chat_id": chat_id, "text": msg, "parse_mode": "Markdown"}, timeout=10)
        except:
            pass

def enviar_notificacao_telegram(nome, email, pacote):
    """ Envia um alerta ao administrador via Telegram quando uma compra é feita """
    token = st.secrets.get("TELEGRAM_BOT_TOKEN")
    chat_id = st.secrets.get("TELEGRAM_ADMIN_CHAT_ID")
    if token and chat_id:
        msg = f"🔔 *Nova Compra de Créditos!*\n\n👤 *Usuário:* {nome}\n📧 *E-mail:* {email}\n📦 *Pacote:* {pacote} Fichas\n\n📌 _Verifique o Painel Admin para realizar a homologação._"
        try:
            requests.post(f"https://api.telegram.org/bot{token}/sendMessage", json={"chat_id": chat_id, "text": msg, "parse_mode": "Markdown"}, timeout=10)
        except:
            pass

def api_obter_produtividade(email):
    """ Busca as linhas de produção do usuário na planilha """
    payload = {"action": "obter_produtividade", "email": email.strip().lower()}
    try:
        resp = requests.post(URL_API_GOOGLE, json=payload, timeout=15)
        res_json = resp.json()
        if res_json.get("success", False):
            return res_json.get("data", [])
        return []
    except:
        return []

def cadastrar_novo_usuario(nome, email, senha, perfil):
    payload = {
        "action": "cadastro", "nome": nome.strip(), "email": email.strip().lower(),
        "senha": hash_senha(senha), "perfil": perfil, "data_cadastro": datetime.now(FUSO_BR).strftime("%d/%m/%Y %H:%M:%S")
    }
    try:
        resp = requests.post(URL_API_GOOGLE, json=payload, timeout=15)
        return resp.json().get("success", False), resp.json().get("message", "Erro do servidor.")
    except Exception as e:
        return False, f"Falha de comunicação: {e}"

def subtrair_credito_nuvem(email, titulo="Não informado", assunto="Não informado"):
    payload = {
        "action": "subtrair_credito", 
        "email": email,
        "titulo": titulo,
        "assunto": assunto
    }
    try:
        resp = requests.post(URL_API_GOOGLE, json=payload, timeout=15)
        res_json = resp.json()
        return res_json.get("success", False), res_json.get("creditos", 0)
    except:
        return False, 0

def enviar_comprovante_nuvem(nome, email, pacote, foto_file):
    img_bytes = foto_file.getvalue()
    img_base64 = base64.b64encode(img_bytes).decode("utf-8")
    
    payload = {
        "action": "enviar_comprovante",
        "nome": nome, "email": email, "pacote": pacote,
        "image_base64": img_base64, "image_type": foto_file.type,
        "data": datetime.now(FUSO_BR).strftime("%d/%m/%Y %H:%M:%S")
    }
    try:
        resp = requests.post(URL_API_GOOGLE, json=payload, timeout=20)
        return resp.json().get("success", False), resp.json().get("message", "Erro.")
    except Exception as e:
        return False, str(e)

def api_obter_pedidos_pendentes():
    try:
        resp = requests.get(URL_API_GOOGLE, timeout=10)
        return resp.json()
    except:
        return []

def aprovar_pedido_nuvem(id_pedido):
    payload = {"action": "aprovar_pedido", "id_pedido": id_pedido}
    try:
        resp = requests.post(URL_API_GOOGLE, json=payload, timeout=15)
        return resp.json().get("success", False)
    except:
        return False

# =====================================================================
# FUNÇÕES ORIGINAIS DO GERADOR DE FICHAS CATALOGRÁFICAS
# =====================================================================
def traduzir_para_portugues(texto):
    try:
        return GoogleTranslator(source='en', target='pt').translate(texto)
    except:
        return texto

def formatar_entrada_aacr2(autores, titulo_obra):
    """
    Formata a entrada principal da ficha seguindo a AACR2:
    - Até 3 autores: Entrada pelo sobrenome do primeiro autor.
    - 4 ou mais autores: Entrada direto pelo título da obra.
    """
    # Garante que temos uma lista de autores limpa e sem espaços vazios
    lista_autores = [a.strip() for a in autores if a.strip()]
    qtd_autores = len(lista_autores)
    
    # REGRA: 4 ou mais autores (Entrada pelo Título)
    if qtd_autores >= 4:
        partes_titulo = titulo_obra.strip().split()
        if partes_titulo:
            # Coloca apenas a primeira palavra do título em CAIXA ALTA
            primeira_palavra = partes_titulo[0].upper()
            resto_titulo = " ".join(partes_titulo[1:])
            return f"{primeira_palavra} {resto_titulo}."
        return f"{titulo_obra.upper()}."
        
    # REGRA: Até 3 autores (Entrada pelo primeiro autor, igual você já fazia)
    elif qtd_autores > 0:
        primeiro_autor = lista_autores[0]
        partes = primeiro_autor.split()
        if len(partes) > 1:
            return f"{partes[-1].upper()}, {' '.join(partes[:-1])}."
        return f"{primeiro_autor.upper()}."
        
    # Caso não venha nenhum autor informado
    return f"{titulo_obra.upper()}."

def remover_artigos(titulo):
    if not titulo: return ""
    artigos = ["O ", "A ", "OS ", "AS ", "UM ", "UMA ", "THE ", "AN "]
    for art in artigos:
        if titulo.upper().startswith(art):
            return titulo[len(art):]
    return titulo

def calcular_cutter(nome_autor):
    try:
        df = pd.read_csv("cutter.csv")
        sobrenome = nome_autor.strip().split()[-1].upper()
        for i in range(len(sobrenome), 2, -1):
            tentativa = sobrenome[:i]
            res = df[df["Name"].str.upper() == tentativa]
            if not res.empty: return str(res.iloc[0]["ID"])
        return "????"
    except: 
        return "????"

@st.cache_data(ttl=3600)
def buscar_descritores_mesh(termo, limite=5):
    url_lookup = "https://id.nlm.nih.gov/mesh/lookup/descriptor"
    params = {"label": termo.strip(), "match": "contains", "limit": limite}
    headers = {"User-Agent": "BiblioKhanMedicas/1.0 (Contato: seu-email@exemplo.com)"}
    
    try:
        resp = requests.get(url_lookup, params=params, headers=headers, timeout=10)
        if resp.status_code != 200 or not resp.json():
            return []

        resultados_completos = []
        for item in resp.json():
            descriptor_id = item.get('resource', '').split('/')[-1]
            termo_oficial = item.get('label', termo) 

            url_details = f"https://id.nlm.nih.gov/mesh/lookup/details?descriptor={descriptor_id}"
            resp_details = requests.get(url_details, headers=headers, timeout=10)
            
            sinonimos_encontrados = []
            if resp_details.status_code == 200:
                data = resp_details.json()
                termos_brutos = data.get('terms', []) + data.get('entryTerms', [])
                for t in termos_brutos:
                    s = t.get('label') or t.get('term') if isinstance(t, dict) else t
                    if s and isinstance(s, str) and s.lower() != termo_oficial.lower():
                        sinonimos_encontrados.append(s)
                sinonimos_encontrados = list(dict.fromkeys(sinonimos_encontrados))
                
            resultados_completos.append({
                "termo_oficial": termo_oficial,
                "sinonimos": sinonimos_encontrados
            })
        return resultados_completos
    except:
        return []

def get_ficha_data(titulo, autores, colaboradores, lista_assuntos, orientador="", coorientador=""):
    autores_v = [a for a in autores if a.strip()]
    entrada = formatar_entrada_autor(autores_v[0]) if autores_v else "AUTOR NÃO INFORMADO"
    sobrenome_letra = autores_v[0].split()[-1][0].upper() if autores_v else "A"
    cutter_id = calcular_cutter(autores_v[0]) if autores_v else "000"
    
    titulo_limpo = remover_artigos(titulo)
    primeira_letra_titulo = titulo_limpo[0].lower() if len(titulo_limpo) > 0 else "a"
    classificacao_cutter = f"{sobrenome_letra}{cutter_id}{primeira_letra_titulo}"
    
    assuntos_limpos = [a for a in lista_assuntos if isinstance(a, str) and a.strip()]
    assuntos = [f"{i+1}. {a.strip().capitalize()}." for i, a in enumerate(dict.fromkeys(assuntos_limpos))]
    
    romanos = ["I", "II", "III", "IV", "V", "VI", "VII", "VIII"]
    r_idx = 0
    
    entradas = [f"{romanos[r_idx]}. Título."]
    r_idx += 1
    
    if orientador.strip():
        entradas.append(f"{romanos[r_idx]}. {formatar_entrada_autor(orientador)}, orient.")
        r_idx += 1
        
    if coorientador.strip():
        entradas.append(f"{romanos[r_idx]}. {formatar_entrada_autor(coorientador)}, coorient.")
        r_idx += 1
        
    for colab in colaboradores:
        if colab["nome"].strip():
            entradas.append(f"{romanos[min(r_idx, len(romanos)-1)]}. {formatar_entrada_autor(colab['nome'])} ({colab['tipo']}).")
            r_idx += 1
    
    return entrada, classificacao_cutter, autores_v, assuntos + entradas

# =====================================================================
# TELA DE LOGIN / CADASTRO
# =====================================================================
if not st.session_state.autenticado:
    try:
        st.image("logo_bibliokhan.png", width=180)
    except:
        st.text("")
        
    st.title("BiblioKhan Médicas")
    st.markdown("### BiblioKhan inteligência e automação para bibliotecas")
    st.caption("Contato de Suporte: Bibliokhancontato@gmail.com")
        
    st.divider()
            
    st.subheader("Identifique-se para acessar o sistema")
    
    aba_login, aba_cadastro = st.tabs(["Iniciar Sessão", "Criar Nova Conta"])
    
    with aba_login:
        log_email = st.text_input("E-mail:")
        log_senha = st.text_input("Senha:", type="password")
        
        if st.button("Entrar no Sistema", use_container_width=True):
            if log_email and log_senha:
                sucesso, nome_user, perfil_user, creditos_user = validar_credenciais(log_email, log_senha)
                if sucesso:
                    st.session_state.autenticado = True
                    st.session_state.user_nome = nome_user
                    st.session_state.user_email = log_email.strip().lower()
                    st.session_state.user_perfil = perfil_user
                    st.session_state.user_creditos = creditos_user
                    
                    st.success(f"Olá, {nome_user}! Login efetuado com sucesso.")
                    st.rerun()
                else:
                    st.error("E-mail ou senha incorretos.")
            else:
                st.warning("Por favor, preencha todos os campos.")
                
    with aba_cadastro:
        cad_nome = st.text_input("Nome Completo:")
        cad_email = st.text_input("E-mail de Acesso:")
        cad_senha = st.text_input("Defina uma Senha:", type="password")
        cad_perfil = st.selectbox("Perfil/Cargo:", ["Bibliotecário(a)", "Residente Médico", "Estudante de Medicina"])
        
        if st.button("Finalizar Cadastro", use_container_width=True):
            if cad_nome and cad_email and cad_senha:
                criado, msg = cadastrar_novo_usuario(cad_nome, cad_email, cad_senha, cad_perfil)
                if criado: 
                    st.success(msg)
                else: 
                    st.error(msg)
            else:
                st.warning("Todos os campos são obrigatórios.")
    st.stop()

# =====================================================================
# INTERFACE PRINCIPAL DO APLICATIVO (SESSÃO ATIVA)
# =====================================================================

# --- BARRA LATERAL (SIDEBAR) ---
try:
    st.sidebar.image("logo_bibliokhan.png", use_container_width=True)
except:
    pass

st.sidebar.markdown(f"### Sessão Ativa")
st.sidebar.markdown(f"**Olá, {st.session_state.user_nome}!**")
st.sidebar.markdown(f"**Cargo:** {st.session_state.user_perfil}")

if st.session_state.user_creditos > 0:
    st.sidebar.success(f"Créditos disponíveis: {st.session_state.user_creditos}")
else:
    st.sidebar.error(f"Créditos: {st.session_state.user_creditos} (Acesso Suspenso)")

if st.sidebar.button("Terminar Sessão", use_container_width=True):
    st.session_state.autenticado = False
    st.rerun()

st.sidebar.markdown("---")
st.sidebar.caption("BiblioKhan inteligência e automação para bibliotecas")
st.sidebar.caption("Bibliokhancontato@gmail.com")

# --- CONFIGURAÇÃO DE ABAS ---
abas_disponiveis = ["Gerar Fichas", "Comprar Fichas", "Produtividade"]
if st.session_state.user_perfil == "Administrador(a)":
    abas_disponiveis.append("Painel Admin")

abas = st.tabs(abas_disponiveis)

# ---------------------------------------------------------------------
# ABA 1: GERADOR DE FICHAS CATALOGRÁFICAS
# ---------------------------------------------------------------------
with abas[0]:
    esta_bloqueado = st.session_state.user_creditos <= 0
    if esta_bloqueado:
        st.error("Função Bloqueada: O seu saldo de fichas terminou. Acesse a aba 'Comprar Fichas' para renovar o seu acesso.")

    col_esq, col_dir = st.columns([1.5, 1], gap="large")

    with col_esq:
        st.subheader("Dados da Obra")
        
        c_tit1, c_tit2 = st.columns(2)
        with c_tit1: titulo = st.text_input("Título da obra:", disabled=esta_bloqueado)
        with c_tit2: titulo_original = st.text_input("Título original (se traduzida):", disabled=esta_bloqueado)
        
        c_pub1, c_pub2, c_pub3 = st.columns(3)
        with c_pub1: cidade = st.text_input("Cidade:", disabled=esta_bloqueado)
        with c_pub2: editora = st.text_input("Editora/Instituição de Defesa:", disabled=esta_bloqueado)
        with c_pub3: ano = st.text_input("Ano:", disabled=esta_bloqueado)
        
        c_desc1, c_desc2, c_desc3, c_desc4 = st.columns(4)
        with c_desc1: volumes = st.text_input("Volume/Edição:", disabled=esta_bloqueado)
        with c_desc2: paginas = st.text_input("Páginas (Ex: 150 p.):", disabled=esta_bloqueado)
        with c_desc3: dimensoes = st.text_input("Dimensões (Ex: 23 cm):", value="23 cm", disabled=esta_bloqueado)
        with c_desc3: isbn = st.text_input("ISBN (deixar vazio para teses):", disabled=esta_bloqueado)
        
        c_class1, c_class2 = st.columns(2)
        with c_class1: classe_principal = st.text_input("Classe Principal DDC/CDU (Ex: 610):", disabled=esta_bloqueado)
        with c_class2: classe_nlm = st.text_input("Classificação NLM (Ex: WG 140):", disabled=esta_bloqueado)

        colecao_serie = st.text_input("Coleção ou Série (Opcional):", disabled=esta_bloqueado)

        st.write("### Trabalho Acadêmico (Teses e Dissertações)")
        e_trabalho_academico = st.checkbox("Esta obra é uma Tese, Dissertação ou Monografia de Residência?", disabled=esta_bloqueado)
        
        grau_academico = "Nenhum"
        area_concentracao = ""
        instituicao = ""
        orientador = ""
        coorientador = ""
        
        if e_trabalho_academico:
            c_acad1, c_acad2 = st.columns(2)
            with c_acad1:
                grau_academico = st.selectbox("Grau Acadêmico:", [
                    "Nenhum", "Dissertação (Mestrado)", "Tese (Doutorado)", 
                    "Tese (Livre-Docência)", "Monografia (Residência Médica)", "Monografia (Especialização)"
                ], disabled=esta_bloqueado)
            with c_acad2:
                area_concentracao = st.text_input("Área de Concentração (Ex: Cardiologia):", disabled=esta_bloqueado)
                
            instituicao = st.text_input("Faculdade/Instituição (Ex: Faculdade de Medicina, Universidade de São Paulo):", disabled=esta_bloqueado)
            
            c_ori1, c_ori2 = st.columns(2)
            with c_ori1: orientador = st.text_input("Nome do Orientador(a):", disabled=esta_bloqueado)
            with c_ori2: coorientador = st.text_input("Nome do Coorientador(a) (Opcional):", disabled=esta_bloqueado)

        st.divider()

        col_autores, col_colab = st.columns(2)
        
        with col_autores:
            st.write("### Autores")
            if st.button("Adicionar Autor", use_container_width=True, disabled=esta_bloqueado): 
                st.session_state.autores.append("")
                st.rerun()
            for i, aut in enumerate(st.session_state.autores):
                c1, c2 = st.columns([8, 2])
                with c1: st.session_state.autores[i] = st.text_input(f"Autor {i+1}", value=aut, key=f"aut_{i}", label_visibility="collapsed", disabled=esta_bloqueado)
                with c2:
                    if st.button("Remover", key=f"del_aut_{i}", disabled=esta_bloqueado) and len(st.session_state.autores) > 1:
                        st.session_state.autores.pop(i)
                        st.rerun()

        with col_colab:
            st.write("### Colaboradores Extensão")
            if st.button("Adicionar Colaborador", use_container_width=True, disabled=esta_bloqueado): 
                st.session_state.colaboradores.append({"nome": "", "tipo": "trad."})
                st.rerun()
            for i, colab in enumerate(st.session_state.colaboradores):
                c1, c2, c3 = st.columns([5, 3, 2])
                with c1: colab["nome"] = st.text_input("Nome", value=colab["nome"], key=f"colab_nome_{i}", label_visibility="collapsed", disabled=esta_bloqueado)
                with c2: colab["tipo"] = st.selectbox("Função", ["trad.", "org.", "comp."], key=f"colab_tipo_{i}", label_visibility="collapsed", disabled=esta_bloqueado)
                with c3:
                    if st.button("Remover", key=f"del_colab_{i}", disabled=esta_bloqueado): 
                        st.session_state.colaboradores.pop(i)
                        st.rerun()

    with col_dir:
        st.subheader("Assuntos e Indexação (MeSH)")
        termo_busca = st.text_input("Buscar termo no MeSH para o Assunto:", disabled=esta_bloqueado)
        
        if termo_busca and not esta_bloqueado:
            if termo_busca != st.session_state.ultimo_termo:
                st.session_state.ultimo_termo = termo_busca
                st.session_state.mesh_limite = 5
                
            resultados = buscar_descritores_mesh(termo_busca, st.session_state.mesh_limite)
            
            if resultados:
                st.success(f"Mostrando até {st.session_state.mesh_limite} termos no banco MeSH.")
                opcoes_nomes = [r["termo_oficial"] for r in resultados]
                escolha = st.selectbox("Selecione o termo mais adequado:", opcoes_nomes)
                termo_escolhido = next(r for r in resultados if r["termo_oficial"] == escolha)
                
                st.markdown(f"### Termo Autorizado (En): **{termo_escolhido['termo_oficial']}**")
                
                if termo_escolhido['sinonimos']:
                    with st.expander(f"Ver sinônimos ({len(termo_escolhido['sinonimos'])} encontrados)"):
                        st.write("Estes termos referem-se ao termo selecionado:")
                        for s in termo_escolhido['sinonimos']: st.markdown(f"- {s}")
                
                qualificadores_comuns = [
                    "Nenhum", "anatomia & histologia", "cirurgia", "citologia", "diagnóstico", 
                    "dietoterapia", "efeitos adversos", "enfermagem", "enzimologia", "epidemiologia", 
                    "ética", "etiologia", "farmacologia", "fisiologia", "fisiopatologia", 
                    "genética", "imunologia", "lesões", "metabolismo", "microbiologia", 
                    "mortalidade", "patologia", "prevenção & controle", "psicologia", 
                    "radiografia", "reabilitação", "sangue", "terapia", "transplante", "urina"
                ]
                qualificador_escolhido = st.selectbox("Adicionar qualificador específico (Opcional):", qualificadores_comuns)
                
                col_add, col_mais = st.columns(2)
                with col_add:
                    if st.button("Adicionar como Assunto", use_container_width=True):
                        termo_em_portugues = traduzir_para_portugues(termo_escolhido['termo_oficial']).capitalize()
                        if qualificador_escolhido != "Nenhum":
                            termo_em_portugues = f"{termo_em_portugues} / {qualificador_escolhido}"
                            
                        if termo_em_portugues not in st.session_state.lista_assuntos:
                            st.session_state.lista_assuntos.append(termo_em_portugues)
                        st.rerun()
                with col_mais:
                    if len(resultados) == st.session_state.mesh_limite:
                        if st.button("Buscar mais resultados", use_container_width=True):
                            st.session_state.mesh_limite += 5
                            st.rerun()

        st.write("### Assuntos Selecionados")
        if not st.session_state.lista_assuntos:
            st.caption("Nenhum assunto selecionado ainda.")
        else:
            for i, assunto in enumerate(st.session_state.lista_assuntos):
                c_assunto, c_del_assunto = st.columns([8, 2])
                with c_assunto: st.markdown(f"• {assunto}")
                with c_del_assunto:
                    if st.button("Remover", key=f"del_assunto_{i}", disabled=esta_bloqueado):
                        st.session_state.lista_assuntos.pop(i)
                        st.rerun()

        st.divider()

        st.subheader("Pré-visualização")
        
        # Busca os dados base do formulário
        entrada_base, class_cutter, auts, lista_final = get_ficha_data(
            titulo, st.session_state.autores, st.session_state.colaboradores, st.session_state.lista_assuntos,
            orientador, coorientador
        )

        # ==========================================
        # CORREÇÃO CRUCIAL AACR2 (Regra de Autores)
        # ==========================================
        qtd_autores = len(auts)

        if qtd_autores >= 4:
            # 1. Entrada Principal vira o TÍTULO (Primeira palavra em CAIXA ALTA)
            partes_titulo = titulo.strip().split()
            if partes_titulo:
                primeira_palavra = partes_titulo[0].upper()
                resto_titulo = " ".join(partes_titulo[1:])
                entrada = f"{primeira_palavra} {resto_titulo}"
            else:
                entrada = titulo.upper()
            
            # 2. Indicação de responsabilidade após a barra (Primeiro autor ... [et al.])
            autores_str = f"{auts[0]} ... [et al.]" if len(auts) > 0 else ""

        else:
            # Mantém a regra padrão para até 3 autores (Entrada pelo sobrenome)
            entrada = entrada_base
            autores_str = ', '.join(auts) if len(auts) > 0 else ''
        # ==========================================

        volumes_str = f"{volumes} ; " if volumes else ""

        # === NOVA LINHA: Formata as dimensões para o padrão ABNT ===
        dimensoes_str = f" ; {dimensoes}" if dimensoes.strip() else ""

        titulo_original_str = f"\n             Título original: {titulo_original}" if titulo_original else ""
        colecao_str = f" ({colecao_serie})" if colecao_serie else ""

        nota_tese_str = ""
        if grau_academico != "Nenhum":
            area_str = f" em {area_concentracao}" if area_concentracao.strip() else ""
            inst_str = f" – {instituicao}" if instituicao.strip() else ""
            nota_tese_str = f"\n             {grau_academico}{area_str}{inst_str}, {cidade}, {ano}."

        bloco_classificacao = []
        if classe_nlm.strip(): bloco_classificacao.append(classe_nlm.strip())
        if classe_principal.strip(): bloco_classificacao.append(classe_principal.strip())

        linhas_class_str = "\n".join(bloco_classificacao)
        bloco_esquerdo_top = f"{linhas_class_str}\n{class_cutter}" if linhas_class_str else class_cutter

        # === ATUALIZADO: {dimensoes_str} inserido logo após {paginas} ===
        ficha_texto = f"""{bloco_esquerdo_top}       {entrada}.
             {titulo} / {autores_str}. – {cidade} : {editora}, {ano}.
             {volumes_str}{paginas}{dimensoes_str}.{colecao_str}{nota_tese_str}{titulo_original_str}
             ISBN {isbn if isbn else "..."}

             {' '.join(lista_final)}"""

        # Código corrigido com aspas triplas para aceitar quebras de linha com segurança
        st.markdown(f"""```text
{ficha_texto}
```""")

        # === BLCO DOS BOTÕES COM A INDENTAÇÃO CORRIGIDA ===
        col_lote_add, col_lote_del = st.columns(2)
        with col_lote_add:
            desativar_lote = esta_bloqueado or not titulo
            if st.button("Adicionar ao Lote", use_container_width=True, disabled=desativar_lote):
                with st.spinner("Processando transação de créditos..."):
                    # Pega a lista do session_state, limpa os espaços e ignora textos vazios
                    if "lista_assuntos" in st.session_state and st.session_state.lista_assuntos:
                        assuntos_limpos = [str(a).strip() for a in st.session_state.lista_assuntos if str(a).strip()]
                        assunto_para_enviar = ", ".join(assuntos_limpos)
                    else:
                        assunto_para_enviar = "Não informado"

                    sucesso_desconto, novos_creditos = subtrair_credito_nuvem(
                        st.session_state.user_email, 
                        titulo=titulo,                  
                        assunto=assunto_para_enviar     
                    )
                    
                    if sucesso_desconto:
                        st.session_state.user_creditos = novos_creditos
                        st.session_state.fichas_lote.append({
                            "classe_nlm": classe_nlm.strip(), "classe_principal": classe_principal.strip(), 
                            "class_cutter": class_cutter, "entrada": entrada, "titulo": titulo, "autores_str": autores_str, 
                            "cidade": cidade, "editora": editora, "ano": ano, "volumes_str": volumes_str, "paginas": paginas, 
                            "colecao_str": colecao_str, "nota_tese_str": nota_tese_str.strip(), "titulo_original_str": titulo_original_str, "isbn": isbn, "lista_final": lista_final
                        })
                        st.success(f"Ficha salva com sucesso. Saldo atualizado. Total no lote: {len(st.session_state.fichas_lote)}")
                        st.rerun()
                    else:
                        st.error("Falha ao validar os créditos na nuvem. Ação interrompida.")
        with col_lote_del:
            if st.button("Limpar Todo o Lote", use_container_width=True):
                st.session_state.fichas_lote = []
                st.rerun()

       if st.session_state.fichas_lote:
            doc = Document()
            for idx, f in enumerate(st.session_state.fichas_lote):
                table = doc.add_table(rows=1, cols=1)
                table.style = 'Table Grid'
                table.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell = table.cell(0, 0)
                table.columns[0].width = Inches(5.3)
                cell.width = Inches(5.3)
                
                p_topo = cell.paragraphs[0]
                p_topo.paragraph_format.space_after = Pt(0)
                p_topo.paragraph_format.line_spacing = 1.15
                
                classes_linhas = []
                if f["classe_nlm"]: classes_linhas.append(f["classe_nlm"])
                if f["classe_principal"]: classes_linhas.append(f["classe_principal"])
                
                if classes_linhas:
                    r_classes = p_topo.add_run("\n".join(classes_linhas))
                    r_classes.font.name = 'Arial'; r_classes.font.size = Pt(10); r_classes.bold = True
                    p_cutter_entrada = cell.add_paragraph()
                else:
                    p_cutter_entrada = p_topo
                    
                p_cutter_entrada.paragraph_format.space_after = Pt(0)
                p_cutter_entrada.paragraph_format.line_spacing = 1.15
                p_cutter_entrada.paragraph_format.tab_stops.add_tab_stop(Inches(0.7)) 
                
                r_cutter = p_cutter_entrada.add_run(f["class_cutter"])
                r_cutter.font.name = 'Arial'; r_cutter.font.size = Pt(10); r_cutter.bold = True
                p_cutter_entrada.add_run("\t") 

                # =========================================================
                # APLICAÇÃO DA REGRA AACR2 NO WORD (4 OU MAIS AUTORES)
                # =========================================================
                # Se houver 'et al.' na string tratada, significa que são 4 ou mais autores
                if 'et al.' in f['autores_str']:
                    # 1. A entrada principal vira o título com a primeira palavra em CAIXA ALTA
                    partes_titulo = f['titulo'].strip().split()
                    if partes_titulo:
                        primeira_palavra = partes_titulo[0].upper()
                        resto_titulo = " ".join(partes_titulo[1:])
                        entrada_docx = f"{primeira_palavra} {resto_titulo}"
                    else:
                        entrada_docx = f['titulo'].upper()
                        
                    # 2. Ajusta de 'Autor et al.' para o padrão oficial 'Autor ... [et al.]'
                    autores_docx_str = f['autores_str']
                    if '...' not in autores_docx_str:
                        autores_docx_str = autores_docx_str.replace("et al.", "... [et al.]")
                else:
                    # Até 3 autores mantém o comportamento normal já gerado
                    entrada_docx = f['entrada']
                    autores_docx_str = f['autores_str']
                # =========================================================
                
                # Aplica a entrada definida pela regra da AACR2
                r_ent = p_cutter_entrada.add_run(f"{entrada_docx}.")
                r_ent.font.name = 'Arial'; r_ent.font.size = Pt(10)
                
                p_corpo = cell.add_paragraph()
                p_corpo.paragraph_format.space_after = Pt(0)
                p_corpo.paragraph_format.line_spacing = 1.15
                p_corpo.paragraph_format.left_indent = Inches(0.7) 
                
                #=== LÓGICA DAS DIMENSÕES ===
                dimensoes_f_str = f" ; {dimensoes}" if dimensoes.strip() else ""

                corpo_linhas = [
                    f"{f['titulo']} / {autores_docx_str}. – {f['cidade']} : {f['editora']}, {f['ano']}.",
                    f"{f['volumes_str']}{f['paginas']}{dimensoes_f_str}.{f['colecao_str']}"
                ]
                if f["nota_tese_str"]:
                    corpo_linhas.append(f["nota_tese_str"])
                if f["titulo_original_str"].strip():
                    corpo_linhas.append(f["titulo_original_str"].strip())
                    
                corpo_linhas.append(f"ISBN {f['isbn'] if f['isbn'] else '...'}")
                corpo_linhas.append("")
                corpo_linhas.append(' '.join(f['lista_final']))
                
                r_corpo = p_corpo.add_run("\n".join(corpo_linhas))
                r_corpo.font.name = 'Arial'; r_corpo.font.size = Pt(10)
                
                if idx < len(st.session_state.fichas_lote) - 1:
                    doc.add_page_break()
                        
            bio = io.BytesIO()
            doc.save(bio)
            
            st.download_button(
                f"Baixar Lote ({len(st.session_state.fichas_lote)} Fichas) (.docx)", 
                data=bio.getvalue(), 
                file_name="lote_fichas_catalograficas.docx", 
                use_container_width=True
            )

# ---------------------------------------------------------------------
# ABA 2: COMPRA DE FICHAS
# ---------------------------------------------------------------------
with abas[1]:
    st.title("Central de Créditos")
    st.write(f"Seu saldo atual: {st.session_state.user_creditos} créditos.")
    
    col_p, col_c = st.columns([1.1, 1])
    with col_p:
        st.markdown("### Pacotes Avulsos")
        st.info("""
        Tabela de valores para recarga de saldo no system:
        * **20 Fichas:** R$ 55,00
        * **30 Fichas:** R$ 80,00
        * **100 Fichas:** R$ 240,00
        * **300 Fichas:** R$ 660,00
        * **600 Fichas:** R$ 1.200,00
        * **800 Fichas:** R$ 1.440,00
        
        ---
        Chave PIX para transferência:
        `bibliokhancontato@gmail.com`
        """)
        
    with col_c:
        st.markdown("### Enviar Comprovante")
        
        opcoes_pacotes = {
            20: "20 Fichas — R$ 55,00",
            30: "30 Fichas — R$ 80,00",
            100: "100 Fichas — R$ 240,00",
            300: "300 Fichas — R$ 660,00",
            600: "600 Fichas — R$ 1.200,00",
            800: "800 Fichas — R$ 1.440,00"
        }
        
        pacote_qtd = st.selectbox(
            "Selecione o pacote adquirido:", 
            options=list(opcoes_pacotes.keys()), 
            format_func=lambda x: opcoes_pacotes[x]
        )
        
        comprovante_file = st.file_uploader("Upload do comprovante Pix (PDF ou Imagem):", type=["png", "jpg", "jpeg"])
        
        if st.button("Enviar para Aprovação", use_container_width=True):
            if comprovante_file:
                with st.spinner("Enviando dados do comprovante..."):
                    ok, msg = enviar_comprovante_nuvem(
                        st.session_state.user_nome, 
                        st.session_state.user_email, 
                        pacote_qtd, 
                        comprovante_file
                    )
                    if ok: 
                        st.success("Comprovante enviado com sucesso. O saldo será atualizado após validação interna.")
                        enviar_notificacao_telegram(st.session_state.user_nome, st.session_state.user_email, pacote_qtd)
                    else: 
                        st.error(f"Erro ao registrar envio: {msg}")
            else: 
                st.warning("Por favor, anexe o comprovante antes de prosseguir.")

# ---------------------------------------------------------------------
# ABA 3: PAINEL DE PRODUTIVIDADE (COMPLETO)
# ---------------------------------------------------------------------
with abas[2]:
    st.title("📊 Painel de Produtividade")
    st.subheader(f"Análise de Indexações de {st.session_state.user_nome}")

    with st.spinner("Carregando dados de produtividade..."):
        dados = api_obter_produtividade(st.session_state.user_email)

    if not dados:
        st.info("Você ainda não possui registros de fichas geradas no lote para criar o gráfico.")
    else:
        import pandas as pd

        # 1. Converte os dados recebidos da API para um DataFrame do Pandas
        df = pd.DataFrame(dados)

        # 2. Coleta todos os assuntos, quebra pelas vírgulas e limpa os espaços
        todos_assuntos = []
        for linha_assunto in df['assunto']:
            if linha_assunto and linha_assunto != "Não informado":
                partes = [a.strip().title() for a in str(linha_assunto).split(",") if a.strip()]
                todos_assuntos.extend(partes)

        # 3. Conta a frequência de cada assunto individual
        if todos_assuntos:
            df_contagem = pd.DataFrame(todos_assuntos, columns=["Assunto"]).value_counts().reset_index(name="Quantidade")
        else:
            df_contagem = pd.DataFrame()

        # 4. Mostra os cartões de resumo (Métricas)
        col_card1, col_card2 = st.columns(2)
        with col_card1:
            st.metric("Total de Livros Processados", len(df))
        with col_card2:
            st.metric("Total de Assuntos Mapeados", len(df_contagem))

        st.markdown("---")
        
        # 5. Renderiza o Gráfico de Barras se houver assuntos mapeados
        if not df_contagem.empty:
            st.write("### 🔝 Assuntos Mais Indexados nas suas Fichas")
            st.bar_chart(
                data=df_contagem,
                x="Assunto",
                y="Quantidade",
                color="#9B5DE5", 
                use_container_width=True
            )
            st.markdown("---")

        # 6. Histórico de Livros Processados e Opção de Download
        st.write("### 📚 Histórico de Obras Processadas")
        
        # Criamos uma cópia limpa para formatar a exibição do usuário
        df_exibicao = df.copy()
        
        # Renomeia as colunas internamente para o relatório ficar amigável
        df_exibicao = df_exibicao.rename(columns={
            "data": "Data/Hora",
            "titulo": "Título da Obra",
            "assunto": "Assuntos Indexados"
        })
        
        # Ajusta a formatação da data para o padrão brasileiro (DD/MM/AAAA HH:MM)
        if "Data/Hora" in df_exibicao.columns:
            try:
                df_exibicao["Data/Hora"] = pd.to_datetime(df_exibicao["Data/Hora"]).dt.strftime('%d/%m/%Y %H:%M')
            except:
                pass 

        # Filtra e organiza apenas as colunas que interessam
        colunas_relatorio = ["Data/Hora", "Título da Obra", "Assuntos Indexados"]
        df_final = df_exibicao[colunas_relatorio]

        # === BOTÃO DE DOWNLOAD ===
        # O encoding 'utf-8-sig' e o separador ';' garantem que o Excel abra certinho no Brasil
        csv_dados = df_final.to_csv(index=False, sep=";").encode('utf-8-sig')
        
        st.download_button(
            label="📥 Baixar Relatório Completo em CSV (Excel)",
            data=csv_dados,
            file_name=f"produtividade_{st.session_state.user_nome.lower().replace(' ', '_')}.csv",
            mime="text/csv",
            use_container_width=True
        )
        
        st.write("") # Pequeno espaçamento visual
        
        # Exibe a tabela visual dinâmica do Streamlit
        st.dataframe(
            df_final, 
            use_container_width=True,
            hide_index=True
        )


# ---------------------------------------------------------------------
# ABA 4: PAINEL ADMIN (CÓDIGO QUE VOCÊ JÁ TINHA, SÓ MUDOU O NÚMERO PARA 3)
# ---------------------------------------------------------------------
if st.session_state.user_perfil == "Administrador(a)":
    with abas[3]: # <-- Mudamos de 2 para 3 aqui!
        st.title("Painel de Approvação de Créditos")
        st.write("Solicitações aguardando processamento:")
        
        pedidos_pendentes = api_obter_pedidos_pendentes()
            
        if not pedidos_pendentes:
            st.info("Nenhuma solicitação pendente de homologação.")
        else:
            for ped in pedidos_pendentes:
                with st.container(border=True):
                    col1, col2 = st.columns([2, 1])
                    with col1:
                        st.markdown(f"**Pedido:** `{ped['id_pedido']}` | **Data:** {ped['data']}")
                        st.markdown(f"**Usuário:** {ped['nome']} ({ped['email']})")
                        st.markdown(f"**Volume Solicitado:** `{ped['pacote']}` Fichas")
                        st.markdown(f"[Visualizar Imagem do Comprovante]({ped['url']})")
                    
                    with col2:
                        st.write("")
                        if st.button(f"Aprovar Créditos", key=f"btn_{ped['id_pedido']}", use_container_width=True):
                            with st.spinner("Atualizando registros no banco de dados..."):
                                if aprovar_pedido_nuvem(ped['id_pedido']):
                                    st.success("Créditos liberados com sucesso.")
                                    st.rerun()
                                else:
                                    st.error("Erro ao tentar atualizar o saldo.")
